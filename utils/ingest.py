import os
import sys
from pathlib import Path
from typing import List
import json
import unstructured_client
from unstructured_client.models import operations, shared
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Add the project root to Python path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from core.storage.database import DocumentStorage, Document
from core.config_loader import config


class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.get("processing.chunk_size", 1000),
            chunk_overlap=config.get("processing.chunk_overlap", 200),
            length_function=len,
        )
        self.unstructured_client = unstructured_client.UnstructuredClient(
            api_key_auth=config.unstructured_api_key
        )

    def process_document(self, file_path: str) -> List[Document]:
        file_path = Path(file_path)
        if not file_path.exists():
            print(f"File not found: {file_path}")
            return []

        print(f"Processing: {file_path.name}")

        text = ""
        element_metadata = []

        if file_path.suffix.lower() == ".pdf":
            text, element_metadata = self._extract_pdf_with_unstructured(file_path)
        elif file_path.suffix.lower() == ".docx":
            text = self._extract_docx_text(file_path)
        elif file_path.suffix.lower() in [".txt", ".md"]:
            text = self._extract_text_file(file_path)
        else:
            print(f"Unsupported file type: {file_path.suffix}")
            return []

        if not text.strip():
            print(f"No text extracted from {file_path.name}")
            return []

        chunks = self.text_splitter.split_text(text)

        documents = []
        for i, chunk in enumerate(chunks):
            if chunk.strip():
                metadata = {
                    "source": str(file_path),
                    "filename": file_path.name,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "file_type": file_path.suffix.lower(),
                }

                if element_metadata and i < len(element_metadata):
                    metadata.update(element_metadata[i])

                documents.append(Document(page_content=chunk, metadata=metadata))

        print(f"Created {len(documents)} chunks from {file_path.name}")
        return documents

    def _extract_pdf_with_unstructured(self, file_path: Path) -> tuple:
        try:
            print(f"Using Unstructured API to process {file_path.name}")

            req = operations.PartitionRequest(
                partition_parameters=shared.PartitionParameters(
                    files=shared.Files(
                        content=open(file_path, "rb"),
                        file_name=str(file_path.name),
                    ),
                    strategy=shared.Strategy.HI_RES,
                    languages=["eng"],
                    split_pdf_page=True,
                    split_pdf_allow_failed=True,
                    split_pdf_concurrency_level=8,
                ),
            )

            res = self.unstructured_client.general.partition(request=req)

            if res.elements:
                text_parts = []
                element_metadata = []

                for element in res.elements:
                    element_dict = dict(element)
                    text_content = element_dict.get("text", "")

                    if text_content and text_content.strip():
                        text_parts.append(text_content)

                        metadata = {
                            "element_type": element_dict.get("type", ""),
                            "page_number": element_dict.get("metadata", {}).get(
                                "page_number", 0
                            ),
                        }

                        coords = element_dict.get("metadata", {}).get("coordinates", {})
                        if coords:
                            metadata["coordinates"] = json.dumps(coords)

                        element_metadata.append(metadata)

                combined_text = "\n\n".join(text_parts)
                print(f"Extracted {len(text_parts)} elements from {file_path.name}")
                return combined_text, element_metadata
            else:
                print(f"No elements extracted from {file_path.name}")
                return "", []

        except Exception as e:
            print(f"Error with Unstructured API for {file_path.name}: {e}")
            print("API extraction failed")
            return "", []

    def _extract_docx_text(self, file_path: Path) -> str:
        try:
            doc = DocxDocument(str(file_path))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            print(f"Error extracting DOCX text: {e}")
            return ""

    def _extract_text_file(self, file_path: Path) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            print(f"Error reading text file: {e}")
            return ""


def process_documents_from_directory(directory: Path = None) -> List[Document]:
    directory = directory or config.docs_dir
    processor = DocumentProcessor()

    all_documents = []

    supported_extensions = [".pdf", ".docx", ".txt", ".md"]

    for file_path in directory.rglob("*"):
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            documents = processor.process_document(file_path)
            all_documents.extend(documents)

    return all_documents


def main():
    print("Document Q&A RAG Bot - Processing Documents")
    print(f"Processing documents from: {config.docs_dir}")

    config.docs_dir.mkdir(exist_ok=True)
    config.data_dir.mkdir(exist_ok=True)

    documents = process_documents_from_directory()

    if not documents:
        print(f"No documents found in {config.docs_dir} folder.")
        print("Add documents and run again.")
        return

    storage = DocumentStorage()
    storage.store_documents(documents)

    print(f"Total documents in database: {storage.get_count()}")
    print("Ready for RAG queries!")

if __name__ == "__main__":
    main()
